from docx import Document
import os

# Load existing document
doc_path = "SSO_Platform_Complete_Documentation.docx"
doc = Document(doc_path)

# Add App3 section
doc.add_heading('Application 3 Integration', level=1)

# Add credentials section
doc.add_heading('Active Directory Credentials', level=2)
doc.add_paragraph('Created AD Users:')

# Test user
p1 = doc.add_paragraph()
p1.add_run('• Username: ').bold = True
p1.add_run('app3user@company.com')
p1 = doc.add_paragraph()
p1.add_run('• Password: ').bold = True  
p1.add_run('App3Pass123!')
p1 = doc.add_paragraph()
p1.add_run('• Role: ').bold = True
p1.add_run('app3-user')

doc.add_paragraph()

# Admin user
p2 = doc.add_paragraph()
p2.add_run('• Username: ').bold = True
p2.add_run('app3admin@company.com')
p2 = doc.add_paragraph()
p2.add_run('• Password: ').bold = True
p2.add_run('App3Admin123!')
p2 = doc.add_paragraph()
p2.add_run('• Role: ').bold = True
p2.add_run('admin, app3-user')

# Add configuration
doc.add_heading('App3 Configuration', level=2)
config_text = """
Port: 5103
Client ID: app3-client
AD Group: App3Users
Keycloak Role Mapping: App3Users → app3-user
"""
doc.add_paragraph(config_text)

# Add SSO flow
doc.add_heading('SSO Authentication Flow', level=2)
flow_text = """
1. User accesses http://localhost:5103
2. Silent authentication check via Keycloak
3. If no session: Full authentication flow
4. LDAP validation against Active Directory
5. Group mapping: App3Users → app3-user role
6. Cross-application SSO enabled
"""
doc.add_paragraph(flow_text)

# Save updated document
doc.save("SSO_Platform_Complete_Documentation_Updated.docx")
print("Document updated with App3 information")