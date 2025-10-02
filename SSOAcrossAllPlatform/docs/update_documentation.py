from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def update_sso_documentation():
    # Open existing document
    doc = Document('D:/AMAZONQ/SSOAcrossAllPlatform/SSO_Platform_Complete_Documentation.docx')
    
    # Add new section for Login Credentials
    doc.add_page_break()
    
    # Add heading
    heading = doc.add_heading('10. Login Credentials & Authentication Methods', level=1)
    
    # Add description
    doc.add_paragraph('The SSO Platform supports multiple authentication methods for different use cases:')
    
    # SSO Portal Credentials
    doc.add_heading('SSO Portal (CommonLogin) Credentials', level=2)
    doc.add_paragraph('Primary authentication method for accessing all applications via Single Sign-On:')
    
    sso_table = doc.add_table(rows=3, cols=2)
    sso_table.style = 'Table Grid'
    sso_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    sso_table.rows[0].cells[0].text = 'URL'
    sso_table.rows[0].cells[1].text = 'http://localhost:5000'
    sso_table.rows[1].cells[0].text = 'Username'
    sso_table.rows[1].cells[1].text = 'sandeepkumar1464@gmail.com'
    sso_table.rows[2].cells[0].text = 'Password'
    sso_table.rows[2].cells[1].text = 'Admin_123'
    
    doc.add_paragraph()
    
    # Application 1 Direct Login
    doc.add_heading('Application 1 Direct Login', level=2)
    doc.add_paragraph('Direct authentication to App1 without SSO (fallback method):')
    
    app1_table = doc.add_table(rows=3, cols=2)
    app1_table.style = 'Table Grid'
    app1_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    app1_table.rows[0].cells[0].text = 'URL'
    app1_table.rows[0].cells[1].text = 'http://localhost:5101/Home/AuthenticateApp1'
    app1_table.rows[1].cells[0].text = 'Username'
    app1_table.rows[1].cells[1].text = 'app1user'
    app1_table.rows[2].cells[0].text = 'Password'
    app1_table.rows[2].cells[1].text = 'App1Pass123'
    
    doc.add_paragraph()
    
    # Application 2 Direct Login
    doc.add_heading('Application 2 Direct Login', level=2)
    doc.add_paragraph('Direct authentication to App2 without SSO (fallback method):')
    
    app2_table = doc.add_table(rows=3, cols=2)
    app2_table.style = 'Table Grid'
    app2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    app2_table.rows[0].cells[0].text = 'URL'
    app2_table.rows[0].cells[1].text = 'http://localhost:5102/Home/AuthenticateApp2'
    app2_table.rows[1].cells[0].text = 'Username'
    app2_table.rows[1].cells[1].text = 'app2user'
    app2_table.rows[2].cells[0].text = 'Password'
    app2_table.rows[2].cells[1].text = 'App2Pass123'
    
    doc.add_paragraph()
    
    # Keycloak Admin Access
    doc.add_heading('Keycloak Admin Console', level=2)
    doc.add_paragraph('Administrative access to Keycloak server:')
    
    keycloak_table = doc.add_table(rows=3, cols=2)
    keycloak_table.style = 'Table Grid'
    keycloak_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    keycloak_table.rows[0].cells[0].text = 'URL'
    keycloak_table.rows[0].cells[1].text = 'http://localhost:8080'
    keycloak_table.rows[1].cells[0].text = 'Username'
    keycloak_table.rows[1].cells[1].text = 'admin'
    keycloak_table.rows[2].cells[0].text = 'Password'
    keycloak_table.rows[2].cells[1].text = 'Admin_123'
    
    doc.add_paragraph()
    
    # Client Secrets Section
    doc.add_heading('Keycloak Client Secrets', level=2)
    doc.add_paragraph('Client secrets for application authentication with Keycloak:')
    
    secrets_table = doc.add_table(rows=4, cols=3)
    secrets_table.style = 'Table Grid'
    secrets_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    secrets_table.rows[0].cells[0].text = 'Application'
    secrets_table.rows[0].cells[1].text = 'Client ID'
    secrets_table.rows[0].cells[2].text = 'Client Secret'
    
    # Data
    secrets_table.rows[1].cells[0].text = 'CommonLogin'
    secrets_table.rows[1].cells[1].text = 'common-login'
    secrets_table.rows[1].cells[2].text = 'b4tvl5GQRT9oiVOSpWnFf2uQHK07jJhF'
    
    secrets_table.rows[2].cells[0].text = 'Application 1'
    secrets_table.rows[2].cells[1].text = 'app1-client'
    secrets_table.rows[2].cells[2].text = '7xsGjfrgp4FjkKV0JcewMgECEKSXYft4'
    
    secrets_table.rows[3].cells[0].text = 'Application 2'
    secrets_table.rows[3].cells[1].text = 'app2-client'
    secrets_table.rows[3].cells[2].text = 'J5nyrCgOZQjqcWRSHIrlLDEHVXxZ3wSU'
    
    doc.add_paragraph()
    
    # Authentication Flow Section
    doc.add_heading('Authentication Flow Options', level=2)
    
    flow_options = [
        'SSO Flow: Login via CommonLogin → Automatic access to App1 and App2',
        'Direct App1: Login directly to App1 using app1user credentials',
        'Direct App2: Login directly to App2 using app2user credentials',
        'Mixed Flow: Use SSO for some apps, direct login for others'
    ]
    
    for option in flow_options:
        doc.add_paragraph(option, style='List Bullet')
    
    # Security Notes
    doc.add_heading('Security Notes', level=2)
    
    security_notes = [
        'SSO credentials provide access to all authorized applications',
        'Direct login credentials are application-specific',
        'Client secrets should be kept secure in production',
        'All passwords should be changed in production environment',
        'Enable HTTPS for all communications in production'
    ]
    
    for note in security_notes:
        doc.add_paragraph(note, style='List Bullet')
    
    # Save the updated document
    doc.save('D:/AMAZONQ/SSOAcrossAllPlatform/SSO_Platform_Complete_Documentation.docx')
    print("Documentation updated successfully!")

if __name__ == "__main__":
    update_sso_documentation()