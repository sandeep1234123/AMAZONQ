from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_sso_documentation():
    # Create a new document
    doc = Document()
    
    # Title
    title = doc.add_heading('SSO Platform - Complete Flow Diagram & Presentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('Single Sign-On Implementation with Keycloak')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. System Architecture Overview',
        '2. Authentication Flow Diagram',
        '3. SSO Token Generation Process',
        '4. Application Integration',
        '5. User Session Management',
        '6. Role-Based Access Control',
        '7. Security Implementation',
        '8. Configuration Details',
        '9. Error Handling & Troubleshooting',
        '10. Production Deployment Guide'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # Section 1: System Architecture
    doc.add_heading('1. System Architecture Overview', level=1)
    
    doc.add_paragraph('The SSO Platform consists of four main components:')
    
    # Architecture table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Port'
    hdr_cells[2].text = 'Purpose'
    
    # Data rows
    components = [
        ('Keycloak Server', '8080', 'Identity Provider & Authentication Server'),
        ('CommonLogin Portal', '5000', 'SSO Gateway & User Dashboard'),
        ('Application 1', '5101', 'Business Application with SSO Integration'),
        ('Application 2', '5102', 'Business Application with SSO Integration')
    ]
    
    for i, (comp, port, purpose) in enumerate(components, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = comp
        row_cells[1].text = port
        row_cells[2].text = purpose
    
    doc.add_paragraph()
    
    # Section 2: Authentication Flow
    doc.add_heading('2. Complete Authentication Flow', level=1)
    
    doc.add_heading('Phase 1: Initial Login Process', level=2)
    
    flow_steps = [
        'User navigates to CommonLogin Portal (http://localhost:5000)',
        'User enters email: sandeepkumar1464@gmail.com',
        'CommonLogin redirects to Keycloak authorization endpoint',
        'Keycloak displays login form',
        'User enters password: Admin_123',
        'Keycloak validates credentials and generates authorization code',
        'Keycloak redirects back to CommonLogin with authorization code',
        'CommonLogin exchanges code for access tokens',
        'CommonLogin displays user dashboard with session information'
    ]
    
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_heading('Phase 2: SSO Token Generation', level=2)
    
    sso_steps = [
        'User clicks "Launch App" from CommonLogin dashboard',
        'CommonLogin validates user roles against application requirements',
        'CommonLogin generates secure SSO token with user information',
        'CommonLogin redirects user to target application with SSO token',
        'Application validates SSO token and creates local session',
        'User gains automatic access without re-authentication'
    ]
    
    for i, step in enumerate(sso_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_page_break()
    
    # Section 3: User Session Information
    doc.add_heading('3. User Session Information Display', level=1)
    
    doc.add_paragraph('The CommonLogin dashboard displays comprehensive user session information:')
    
    # Session info table
    session_table = doc.add_table(rows=8, cols=2)
    session_table.style = 'Table Grid'
    
    session_data = [
        ('User ID', 'Keycloak subject identifier (sub claim)'),
        ('Username', 'preferred_username from Keycloak'),
        ('Email', 'User email address'),
        ('First Name', 'given_name claim'),
        ('Last Name', 'family_name claim'),
        ('Roles', 'Assigned roles from Keycloak'),
        ('Session Started', 'Current timestamp'),
        ('SSO Status', 'Active across all applications')
    ]
    
    for i, (field, description) in enumerate(session_data):
        row_cells = session_table.rows[i].cells
        row_cells[0].text = field
        row_cells[1].text = description
    
    # Section 4: Role-Based Access Control
    doc.add_heading('4. Role-Based Access Control Matrix', level=1)
    
    # RBAC table
    rbac_table = doc.add_table(rows=5, cols=5)
    rbac_table.style = 'Table Grid'
    
    # Header
    rbac_headers = ['User Role', 'CommonLogin', 'App1', 'App2', 'Description']
    for i, header in enumerate(rbac_headers):
        rbac_table.rows[0].cells[i].text = header
    
    # RBAC data
    rbac_data = [
        ('admin', '✓', '✓', '✓', 'Full access to all applications'),
        ('app1-user', '✓', '✓', '✗', 'Access to App1 only'),
        ('app2-user', '✓', '✗', '✓', 'Access to App2 only'),
        ('multi-user', '✓', '✓', '✓', 'Access to App1 and App2')
    ]
    
    for i, (role, common, app1, app2, desc) in enumerate(rbac_data, 1):
        row_cells = rbac_table.rows[i].cells
        row_cells[0].text = role
        row_cells[1].text = common
        row_cells[2].text = app1
        row_cells[3].text = app2
        row_cells[4].text = desc
    
    doc.add_page_break()
    
    # Section 5: Security Features
    doc.add_heading('5. Security Implementation', level=1)
    
    doc.add_heading('Token Security', level=2)
    security_features = [
        'Token Expiry: 5-minute lifetime for SSO tokens',
        'Timestamp Validation: Prevents replay attacks',
        'Role Verification: Each application validates user roles',
        'Secure Transport: HTTPS in production environment'
    ]
    
    for feature in security_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Session Management', level=2)
    session_features = [
        'HttpOnly Cookies: Prevent XSS attacks',
        'SameSite Policy: CSRF protection',
        'Secure Flag: HTTPS-only cookies in production',
        'Session Timeout: 8-hour sliding expiration'
    ]
    
    for feature in session_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Section 6: Configuration
    doc.add_heading('6. System Configuration', level=1)
    
    doc.add_heading('Keycloak Configuration', level=2)
    doc.add_paragraph('Realm: sso-realm')
    doc.add_paragraph('Clients:')
    
    clients = [
        'common-login (http://localhost:5000)',
        'app1-client (http://localhost:5101)',
        'app2-client (http://localhost:5102)'
    ]
    
    for client in clients:
        doc.add_paragraph(client, style='List Bullet')
    
    doc.add_heading('Test User Configuration', level=2)
    doc.add_paragraph('Primary User: sandeepkumar1464@gmail.com')
    doc.add_paragraph('Password: Admin_123')
    doc.add_paragraph('Roles: admin, app1-user, app2-user')
    
    doc.add_page_break()
    
    # Section 7: Logout Process
    doc.add_heading('7. Logout Process Flow', level=1)
    
    logout_steps = [
        'User clicks "Logout" button in CommonLogin dashboard',
        'CommonLogin clears all local session cookies',
        'CommonLogin initiates Keycloak logout with id_token_hint',
        'Keycloak terminates global SSO session',
        'All application sessions become invalid',
        'User redirected to login page',
        'Re-authentication required for next access'
    ]
    
    for i, step in enumerate(logout_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    # Section 8: Error Handling
    doc.add_heading('8. Error Handling & Troubleshooting', level=1)
    
    doc.add_heading('Common Issues and Solutions', level=2)
    
    errors = [
        ('Client Not Found', 'Verify client exists in Keycloak sso-realm'),
        ('Invalid Client Secret', 'Update appsettings.json with correct secret'),
        ('Access Denied', 'Check user role assignments in Keycloak'),
        ('Correlation Failed', 'Clear browser cookies and restart session'),
        ('Port Binding Error', 'Use different ports or kill existing processes')
    ]
    
    error_table = doc.add_table(rows=len(errors)+1, cols=2)
    error_table.style = 'Table Grid'
    
    error_table.rows[0].cells[0].text = 'Error'
    error_table.rows[0].cells[1].text = 'Solution'
    
    for i, (error, solution) in enumerate(errors, 1):
        error_table.rows[i].cells[0].text = error
        error_table.rows[i].cells[1].text = solution
    
    # Section 9: Production Deployment
    doc.add_heading('9. Production Deployment Considerations', level=1)
    
    doc.add_heading('Security Enhancements', level=2)
    prod_security = [
        'Enable HTTPS for all communications',
        'Use strong client secrets (32+ characters)',
        'Implement proper JWT signing with RSA keys',
        'Add rate limiting and DDoS protection',
        'Enable comprehensive audit logging'
    ]
    
    for item in prod_security:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Performance Optimizations', level=2)
    prod_performance = [
        'Use Redis for distributed session storage',
        'Implement connection pooling',
        'Enable token caching',
        'Configure load balancing',
        'Set up CDN for static assets'
    ]
    
    for item in prod_performance:
        doc.add_paragraph(item, style='List Bullet')
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('SSO Platform Documentation\nGenerated for AMAZONQ Project\n© 2025')
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    
    # Save the document
    doc_path = 'D:/AMAZONQ/SSOAcrossAllPlatform/SSO_Platform_Complete_Documentation.docx'
    doc.save(doc_path)
    print(f"Documentation created successfully: {doc_path}")

if __name__ == "__main__":
    create_sso_documentation()